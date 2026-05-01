"""DOCX exporter.
"""
from __future__ import annotations

import logging
import os
from io import BytesIO
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.models import FileRef
from app.core.templates import TemplateRegistry

log = logging.getLogger(__name__)
log.debug("DOCX exporter initialized")


def _convert_markdown_to_structured(content: str) -> list[dict]:
    """
    Convert a markdown string to a list of structured blocks.
    :param content: The markdown string to convert
    :return: A list of structured blocks
    """
    return [{"type": "paragraph", "text": line} for line in content.splitlines() if line.strip()]


def _normalize_content(content: list[dict] | str) -> list[dict]:
    """
    Normalize the content to a list of structured blocks.
    :param content: The content to normalize
    :return: A list of structured blocks
    """
    if isinstance(content, str):
        return _convert_markdown_to_structured(content)
    if isinstance(content, list):
        return content
    return []


def _generate_unique_folder() -> str:
    """
    Generate a unique folder.
    :return: A unique folder
    """
    return "./tmp"


def _generate_filename(folder_path: str, extension: str) -> tuple[str, str]:
    """
    Generate a filename.
    :param folder_path: The folder path
    :param extension: The extension
    :return: A tuple of the filename and the folder path
    """
    fname = f"document.{extension}"
    return os.path.join(folder_path, fname), fname


def _public_url(folder_path: str, filename: str) -> str:
    """
    Generate a public URL.
    :param folder_path: The folder path
    :param filename: The filename
    :return: A public URL
    """
    return f"{folder_path}/{filename}"


def _process_docx_template(template_path: str, replacements: dict[str, str]) -> Any:
    """
    Process a DOCX template.
    :param template_path: The path to the template
    :param replacements: The replacements to make
    :return: A DOCX document
    """
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(str(key), str(value))
    return doc


def _load_style_template(template_path: str) -> Any:
    """
    Load a style template.
    :param template_path: The path to the template
    :return: A DOCX document
    """

    doc = Document(template_path)
    for paragraph in doc.paragraphs[:]:
        for run in paragraph.runs:
            run.text = ""
    return doc


def _maybe_apply_title(
    *, 
    doc: Any, 
    title: str | None, 
    use_template: Optional[str], 
    process_as_template: bool
) -> None:
    """
    Maybe apply a title.
    :param doc: The DOCX document
    :param title: The title to apply
    :param use_template: Whether to use a template
    :param process_as_template: Whether to process as a template
    """
    if not title or process_as_template:
        return

    title_added = False
    if use_template:
        for paragraph in doc.paragraphs:
            if paragraph.style.name == "Title" or "Title" in paragraph.style.name:
                paragraph.text = title
                title_added = True
                break

    if not title_added:
        title_paragraph = doc.add_paragraph(title)
        try:
            title_paragraph.style = doc.styles["Title"]
        except KeyError:
            try:
                title_paragraph.style = doc.styles["Heading 1"]
            except KeyError:
                run = title_paragraph.runs[0] if title_paragraph.runs else title_paragraph.add_run()
                run.font.size = 20
                run.font.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _append_content(
    *, 
    doc: Any, 
    content: list[dict], 
    use_template: Optional[str], 
    log: logging.Logger
) -> None:
    """
    Append content to the DOCX document.
    :param doc: The DOCX document
    :param content: The content to append
    :param use_template: Whether to use a template
    :param log: The logger
    """
    last_usable_paragraph = None
    if doc.paragraphs:
        for paragraph in reversed(doc.paragraphs):
            if not paragraph.text.strip():
                last_usable_paragraph = paragraph
                break

    for item in content or []:
        if isinstance(item, str):
            last_usable_paragraph = _render_paragraph_item(
                doc=doc,
                item={"type": "paragraph", "text": item},
                use_template=use_template,
                last_usable_paragraph=last_usable_paragraph,
                log=log,
            )
            continue
        if not isinstance(item, dict):
            continue

        item_type = item.get("type")
        if item_type == "image_query":
            log.warning("image_query is not supported in current DOCX exporter")
            continue

        renderer = {
            "title": _render_title_item,
            "subtitle": _render_subtitle_item,
            "heading": _render_heading_item,
            "paragraph": _render_paragraph_item,
            "list": _render_list_item,
            "table": _render_table_item,
        }.get(item_type, _render_text_fallback_item)

        last_usable_paragraph = renderer(
            doc=doc,
            item=item,
            use_template=use_template,
            last_usable_paragraph=last_usable_paragraph,
            log=log,
        )


def _render_title_item(
    *, 
    doc: Any, 
    item: dict, 
    use_template: Optional[str], 
    last_usable_paragraph: Any, 
    log: logging.Logger
) -> Any:
    """
    Render a title item.
    :param doc: The DOCX document
    :param item: The item to render
    :param use_template: Whether to use a template
    :param last_usable_paragraph: The last usable paragraph
    :param log: The logger
    """
    paragraph = doc.add_paragraph(item.get("text", ""))
    try:
        paragraph.style = doc.styles["Heading 1"]
    except KeyError:
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = 18
        run.font.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return last_usable_paragraph


def _render_subtitle_item(
    *, 
    doc: Any, 
    item: dict, 
    use_template: Optional[str], 
    last_usable_paragraph: Any, 
    log: logging.Logger
) -> Any:
    paragraph = doc.add_paragraph(item.get("text", ""))
    try:
        paragraph.style = doc.styles["Heading 2"]
    except KeyError:
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = 16
        run.font.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return last_usable_paragraph


def _render_heading_item(
    *, 
    doc: Any, 
    item: dict, 
    use_template: Optional[str], 
    last_usable_paragraph: Any, 
    log: logging.Logger
) -> Any:
    """
    """
    level = item.get("level")
    if level == 1:
        return _render_title_item(doc=doc, item=item, use_template=use_template, last_usable_paragraph=last_usable_paragraph, log=log)
    if level == 2:
        return _render_subtitle_item(doc=doc, item=item, use_template=use_template, last_usable_paragraph=last_usable_paragraph, log=log)
    return _render_paragraph_item(doc=doc, item=item, use_template=use_template, last_usable_paragraph=last_usable_paragraph, log=log)


def _apply_paragraph_style_if_needed(
    *, 
    doc: Any, 
    paragraph: Any, 
    use_template: Optional[str]
) -> None:
    if not use_template:
        return
    try:
        paragraph.style = doc.styles["Normal"]
    except Exception:
        pass


def _render_paragraph_item(
    *, 
    doc: Any, 
    item: dict, 
    use_template: Optional[str], 
    last_usable_paragraph: Any, 
    log: logging.Logger
) -> Any:
    if last_usable_paragraph and not last_usable_paragraph.text.strip():
        last_usable_paragraph.text = item.get("text", "")
        return None

    paragraph = doc.add_paragraph(item.get("text", ""))
    _apply_paragraph_style_if_needed(doc=doc, paragraph=paragraph, use_template=use_template)
    return last_usable_paragraph


def _render_list_item(
    *, doc: Any, 
    item: dict, 
    use_template: Optional[str], 
    last_usable_paragraph: Any, 
    log: logging.Logger
) -> Any:
    for item_text in item.get("items", []):
        paragraph = doc.add_paragraph(item_text)
        try:
            paragraph.style = doc.styles["List Bullet"]
        except KeyError:
            try:
                paragraph.style = doc.styles["List Paragraph"]
            except KeyError:
                paragraph.style = doc.styles["Normal"]
    return last_usable_paragraph


def _render_table_item(
    *, doc: Any, 
    item: dict, 
    use_template: Optional[str], 
    last_usable_paragraph: Any, 
    log: logging.Logger
) -> Any:

    data = item.get("data", [])
    if not data:
        return last_usable_paragraph

    table = doc.add_table(rows=len(data), cols=len(data[0]) if data else 0)
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            cell_obj = table.cell(i, j)
            cell_obj.text = str(cell)
            if i == 0:
                for paragraph in cell_obj.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return last_usable_paragraph


def _render_text_fallback_item(
    *, 
    doc: Any, 
    item: dict, 
    use_template: Optional[str], 
    last_usable_paragraph: Any, 
    log: logging.Logger
) -> Any:
    if "text" not in item:
        return last_usable_paragraph
    if last_usable_paragraph and not last_usable_paragraph.text.strip():
        last_usable_paragraph.text = item["text"]
        return None

    paragraph = doc.add_paragraph(item["text"])
    _apply_paragraph_style_if_needed(doc=doc, paragraph=paragraph, use_template=use_template)
    return last_usable_paragraph


def export_docx(
    content: list[dict] | str,
    filename: str | None,
    template_vars: dict[str, str] | None = None,
) -> FileRef:
    """Export *content* as a Word document.

    :param content: Structured content list or markdown string.
    :param filename: Output filename.
    :param template_vars: Placeholder → value mapping for template-fill mode.
    :return: A dictionary with the url and path of the file
    """
    result = _create_word(
        content=content,
        filename=filename,
        folder_path=None,
        title=None,
        use_template=None,
        template_vars=template_vars,
    )
    return FileRef(url=result["url"], path=result["path"])  # type: ignore[arg-type]


def _create_word(
    content: list[dict] | str,
    filename: str,
    folder_path: str | None = None,
    title: str | None = None,
    use_template: Optional[str] = None,
    template_vars: dict = None
) -> dict:
    """Create a Word document.

    :param content: Document content (text or structured data)
    :param filename: File name
    :param folder_path: Save path
    :param title: Document title
    :param use_template: Template name or True for template usage
    :param template_vars: Variables for replacement in template
    :return: A dictionary with the url and path of the file
    """
    log = logging.getLogger(__name__)
    log.debug(f"Creating Word document, use_template={use_template}, template_vars={template_vars}")

    process_as_template = bool(template_vars and isinstance(template_vars, dict))
    replacements: dict[str, str] = template_vars if process_as_template else {}

    if folder_path is None:
        folder_path = _generate_unique_folder()
    if filename:
        filepath = os.path.join(folder_path, filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        fname = filename
    else:
        filepath, fname = _generate_filename(folder_path, "docx")

    template_path = TemplateRegistry.get("docx")
    from docx import Document

    # Mode 1: Template variable filling
    if process_as_template and replacements:
        if template_path:
            doc = _process_docx_template(template_path, replacements)
        else:
            log.warning("TemplateRegistry has no 'docx' template; using empty document")
            doc = Document()
        doc.save(filepath)
        return {"url": _public_url(folder_path, fname), "path": filepath}

    # Mode 2: Style template usage (if requested)
    if use_template and template_path:
        try:
            doc = _load_style_template(template_path)
        except Exception:
            log.exception("Failed to load style template")
            doc = Document()
    else:
        doc = Document()

    _maybe_apply_title(doc=doc, title=title, use_template=use_template, process_as_template=process_as_template)

    # Mode 3: No template (or after style template load), render content
    if not process_as_template:
        structured = _normalize_content(content)
        _append_content(doc=doc, content=structured, use_template=use_template, log=log)

    doc.save(filepath)
    return {"url": _public_url(folder_path, fname), "path": filepath}
