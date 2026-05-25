"""MCP tools: export_text_file, export_document.

Thin wrappers only — no business logic here.
All work is delegated to formats/* modules.
"""
from __future__ import annotations

from typing import Any, Optional
import re
from pathlib import Path
from docx import Document

from app import mcp
from app.core.models import ExportError, ExportFormat
from formats.csv import export_csv
from formats.docx import export_docx
from formats.pdf import export_pdf
from formats.pptx import export_pptx
from formats.raw import export_raw
from formats.xlsx import export_xlsx
from storage.manager import StorageManager


@mcp.tool()
async def export_text_file(
    text: str | list,
    filename: str | None = None,
    format: str = "txt",
) -> dict[str, Any]:
    """Export text or structured content to a file.

    Use this tool when the user asks to save, export, or download content as a file.

    :param format: Target file format. Choose based on content type:
        - "txt"  — plain text, no formatting
        - "md"   — Markdown text
        - "html" — HTML markup string
        - "json" — JSON string (pass serialized JSON as text)
        - "xml"  — XML string
        - "csv"  — tabular data; requires text to be a list of row-lists (see below)
        - "docx" — Word document; requires text to be a list of content-block dicts (see below)
        - "pdf"  — PDF; same structure as docx
        - "pptx" — PowerPoint; same structure as docx
        - "xlsx" — Excel spreadsheet; requires text to be a list of row-lists (see below)

    :param filename: Optional output filename without path (e.g. "report.docx").
        Extension is added automatically if missing. If omitted, a name is generated.

    :param text: Content to export. Type depends on format:

        FOR txt / md / html / json / xml — pass a plain Python string:
            "Hello world"

        FOR csv / xlsx — pass a list of rows, each row is a list of cell values.
            First row is treated as the header:
            [
                ["Name", "Score", "Grade"],
                ["Alice", 95, "A"],
                ["Bob",   88, "B"],
            ]

        FOR docx / pdf / pptx — pass a FLAT list of content-block dicts.
            Do NOT wrap blocks in an outer dict or add a "content" key around them.
            Supported block types:

            {"type": "heading",   "level": 1, "text": "Document title"}  # level 1–6
            {"type": "heading",   "level": 2, "text": "Section heading"}
            {"type": "paragraph", "text": "Body text goes here."}
            {"type": "list",      "items": ["First item", "Second item", "Third item"]}
            {"type": "table",     "data": [
                ["Column A", "Column B"],   # first row = bold header
                ["value 1",  "value 2"],
            ]}

            Full example — a docx with a title, intro, bullet list, and table:
            [
                {"type": "heading",   "level": 1, "text": "Q1 Sales Report"},
                {"type": "paragraph", "text": "This report summarises Q1 performance."},
                {"type": "list",      "items": ["Revenue up 12%", "Churn down 3%"]},
                {"type": "table",     "data": [
                    ["Region", "Revenue"],
                    ["North",  "$120k"],
                    ["South",  "$98k"],
                ]},
            ]

            WRONG — never wrap in an outer object; this produces an empty file:
            [{"title": "Report", "content": [...]}]   <- incorrect

    :returns: {"url": "...", "path": "...", "name": "..."} on success,
              or {"error": {"message": "...", "code": "..."}} on failure.
    """
    try:
        fmt = ExportFormat(format.lower())
    except ValueError:
        return ExportError(
            message=f"Unsupported format: {format}",
            code="UNSUPPORTED_FORMAT"
        ).to_dict()

    try:
        if fmt in ExportFormat.text_formats():
            return export_raw(text, filename, fmt).to_dict()

        elif fmt == ExportFormat.CSV:
            if not isinstance(text, list):
                return ExportError(message="CSV format requires list input").to_dict()
            return export_csv(text, filename).to_dict()

        elif fmt == ExportFormat.DOCX:
            return export_docx(text, filename).to_dict()

        elif fmt == ExportFormat.PPTX:
            return export_pptx(text, filename).to_dict()

        elif fmt == ExportFormat.XLSX:
            if not isinstance(text, list):
                return ExportError(message="XLSX format requires list input").to_dict()
            return export_xlsx(text, filename).to_dict()

        elif fmt == ExportFormat.PDF:
            return (await export_pdf(text, filename)).to_dict()

    except ExportError as exc:
        return exc.to_dict()
    except NotImplementedError as exc:
        return ExportError(message=str(exc), code="NOT_IMPLEMENTED").to_dict()
    except Exception as exc:
        return ExportError(message=str(exc), code="INTERNAL_ERROR").to_dict()


@mcp.tool()
async def export_document(
    data: list[list[str]],
    filename: str | None = None,
    format: str = "csv",
    title: str | None = None,
) -> dict[str, Any]:
    """Export a two-dimensional table to a structured document.

    Use this tool when the content is purely tabular (rows and columns) and the user
    wants csv, xlsx, docx, or pptx output. For rich documents with headings, paragraphs,
    and mixed content, use export_text_file instead.

    :param data: A list of rows; each row is a list of cell values (str or coercible to str).
        The first row is treated as the header / column names in all formats.
        Example:
            [
                ["Name",  "Department", "Salary"],
                ["Alice", "Engineering", "$120k"],
                ["Bob",   "Marketing",   "$95k"],
            ]

    :param filename: Optional output filename without path (e.g. "employees.xlsx").
        Extension is added automatically if missing.

    :param format: Target file format. Supported values:
        - "csv"  — comma-separated values (default)
        - "xlsx" — Excel spreadsheet; pass title to set the sheet name
        - "docx" — Word document with a proper table (first row = bold header)
        - "pptx" — PowerPoint; each row becomes one slide
                   (first cell = slide title, remaining cells = body text)

    :param title: Optional title string.
        For xlsx: used as the worksheet name.
        Ignored for other formats.

    :returns: {"url": "...", "path": "...", "name": "..."} on success,
              or {"error": {"message": "...", "code": "..."}} on failure.
    """
    try:
        fmt = ExportFormat(format.lower())
    except ValueError:
        return ExportError(
            message=f"Unsupported format: {format}",
            code="UNSUPPORTED_FORMAT"
        ).to_dict()

    try:
        if fmt == ExportFormat.CSV:
            return export_csv(data, filename).to_dict()

        elif fmt == ExportFormat.XLSX:
            return export_xlsx(data, filename, title=title).to_dict()

        elif fmt == ExportFormat.DOCX:
            blocks = [{"type": "table", "data": data}]
            return export_docx(blocks, filename).to_dict()

        elif fmt == ExportFormat.PPTX:
            slides = [
                {
                    "title": row[0] if row else "Slide",
                    "content": "\n".join(str(c) for c in row),
                }
                for row in data
            ]
            return export_pptx(slides, filename).to_dict()

        else:
            return ExportError(
                message=f"Unsupported format: {format}",
                code="UNSUPPORTED_FORMAT"
            ).to_dict()

    except ExportError as exc:
        return exc.to_dict()
    except NotImplementedError as exc:
        return ExportError(
            message=str(exc),
            code="NOT_IMPLEMENTED"
        ).to_dict()
    except Exception as exc:
        return ExportError(
            message=str(exc),
            code="INTERNAL_ERROR"
        ).to_dict()


@mcp.tool()
async def fill_document_template(
    template_name: str,
    vars: dict[str, str],
    filename: Optinal[str] = None,
) -> dict[str, Any]:
    """Fill a named .docx template with placeholder values and export the result.

    Templates are .docx files stored in the templates/ directory.
    Placeholders in the document use {{key}} syntax and are replaced
    with the corresponding values from `vars`.

    Use this tool for structured documents: protocols, letters, orders,
    proxies, contracts, and any other named Word templates.

    :param template_name: Name of the template file (without extension),
        e.g. "protocol", "letter", "order", "proxy".
        Use list_document_templates() to see available templates.

    :param vars: Dict of placeholder replacements, e.g.:
        {
            "date": "25.05.2026",
            "recipient": "Goodchild Trevor",
            "amount": "150 000 $"
        }

    :param filename: Optional output filename (e.g. "protocol_2026.docx").
        If omitted, generated automatically.

    :returns: {"url": "...", "path": "...", "name": "..."} on success,
              or {"error": {"message": "...", "code": "..."}} on failure.
    """

    templates_dir = Path(os.getenv("OWN_TEMPLATES_PATH", "templates"))
    template_path = templates_dir / f"{template_name}.docx"

    if not template_path.exists():
        available = [p.stem for p in templates_dir.glob("*.docx")]
        return ExportError(
            message=f"Template '{template_name}' not found. Available: {available}",
            code="TEMPLATE_NOT_FOUND",
        ).to_dict()

    def replace_in_text(text: str) -> str:
        return re.sub(
            r"\{\{(.+?)\}\}",
            lambda m: str(vars.get(m.group(1).strip(), m.group(0))),
            text,
        )

    doc = Document(str(template_path))

    for para in doc.paragraphs:
        for run in para.runs:
            run.text = replace_in_text(run.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = replace_in_text(run.text)

    storage = StorageManager()
    folder_path = storage.generate_folder()
    out_filename = os.path.basename(filename or f"{template_name}_filled.docx")
    if not out_filename.endswith(".docx"):
        out_filename += ".docx"
    out_path = os.path.join(folder_path, out_filename)
    doc.save(out_path)

    from app.core.models import ExportResult
    from storage.manager import public_url
    return ExportResult(
        url=public_url(folder_path, out_filename),
        path=out_path,
        name=out_filename,
    ).to_dict()


@mcp.tool()
def list_document_templates() -> dict[str, Any]:
    """List all available named .docx templates.

    Returns template names that can be passed to fill_document_template().

    :returns: {"templates": ["protocol", "letter", "order", "proxy", ...]}
    """
    from pathlib import Path

    templates_dir = Path(os.getenv("DOCS_TEMPLATE_PATH", "templates"))
    if not templates_dir.exists():
        return {"templates": [], "error": "Templates directory not found"}

    names = sorted(p.stem for p in templates_dir.glob("*.docx")
                   if p.stem != "Default_Template")
    return {"templates": names}
